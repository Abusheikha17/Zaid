from docx import Document
from docx.shared import Inches
import pyttsx3
def speak(text):
    pyttsx3.speak(text)

document = Document()

#Photo
document.add_picture('Zaid.JPG', width= Inches (2.0))


Name = input('What is your name? ')
Phone_number = input('What is your phone number? ')
Email = input('What is your email? ')

document.add_paragraph(
    Name + ' | ' + Phone_number + ' | ' + Email)

#about me
document.add_heading('About me')
about_me = input('Tell about yourself?')
document.add_paragraph(about_me)

#Work Experience 
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company name')
from_date = input('From Date')
to_date = input('To Date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date +'\n').italic =  True

experience_details = input(
    'Decribe your experience at ' + company)
p.add_run(experience_details)

#More Experience
while True:
    has_more_Experience  = input(
        'Do you have more experience? Yes or No ')
    if has_more_Experience.lower() == 'yes' :
        p = document.add_paragraph()

        company = input('Enter company name ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date +'\n').italic =  True

        experience_details = input(
         'Decribe your experience at ' + company )
        p.add_run(experience_details)
    else:
        break    
#Skills 
# Skills
document.add_heading('Skills')
skill = input('Enter Skill ')
p = document.add_paragraph(skill)
p.style = 'ListBullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter Skill ')
        p = document.add_paragraph(skill)
        p.style = 'ListBullet'
    else:
        break



#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Abusheikha for exchange and Uiinted Airlines "

    
       
document.save('cv.docx')
